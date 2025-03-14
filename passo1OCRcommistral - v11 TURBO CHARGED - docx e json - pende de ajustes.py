import traceback
from mistralai import Mistral  # Você precisará instalar esta biblioteca: pip install mistralai
import os
import logging
from base64 import b64encode, b64decode
import re
import pandas as pd
from datetime import datetime
from typing import Dict, List, Any, Optional
import json
import os
import sys
import subprocess
import logging
import tkinter as tk
from tkinter import messagebox
import platform
import json
import requests
import uuid
import base64
import threading
write_lock = threading.Lock()
import datetime
import re
import winreg  # Para verificação do PATH no Windows
from typing import Dict, List, Tuple, Optional, Any
import os
import sys
import subprocess
import logging
import tkinter as tk
from tkinter import messagebox
import platform
import json
import requests
import uuid
import base64

# ------------------------- VERIFICAÇÃO DE DEPENDÊNCIAS -------------------------
# Lista de dependências necessárias
required_modules = {
    "python-docx": "docx",
    "beautifulsoup4": "bs4",
    "pytesseract": "pytesseract",
    "pdf2image": "pdf2image",
    "Pillow": "PIL",
    "pdfminer.six": "pdfminer",
    "pydantic": "pydantic",
    "pandas": "pandas",
    "mistralai": "mistralai"
}

# Verifica e instala módulos ausentes
missing_modules = []
for pip_name, import_name in required_modules.items():
    try:
        __import__(import_name)
    except ImportError:
        missing_modules.append(pip_name)

# Instala módulos faltantes
if missing_modules:
    print(f"Instalando módulos: {', '.join(missing_modules)}")
    for module in missing_modules:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", module])
            print(f"Módulo {module} instalado com sucesso!")
        except subprocess.CalledProcessError:
            print(f"Falha ao instalar {module}. Por favor, instale manualmente.")
            if 'tkinter' in sys.modules:
                root = tk.Tk()
                root.withdraw()
                messagebox.showerror("Erro", f"Não foi possível instalar {module}. Execute: pip install {module}")
                root.destroy()
            exit(1)
    
    # Reinicia o aplicativo após instalar os módulos
    print("Reiniciando aplicativo com as novas dependências...")
    os.execv(sys.executable, [sys.executable] + sys.argv)

# Agora que verificamos as dependências, importamos os demais módulos
import threading
import datetime
from html import escape, unescape
from tkinter import ttk, filedialog, scrolledtext
from typing import Dict, Optional, List, Tuple, Union, Any
from concurrent.futures import ThreadPoolExecutor, as_completed
from pdfminer.high_level import extract_text
from PIL import Image, ImageOps
import pytesseract
import pdf2image
import queue
from bs4 import BeautifulSoup
from logging.handlers import QueueHandler, RotatingFileHandler
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re
import tempfile
from io import BytesIO


# ------------------------- CONSTANTES -------------------------
SUPPORTED_LANGS = ['por', 'eng', 'spa', 'fra', 'deu']
MIN_TEXT_LENGTH = 50
MAX_LOG_SIZE = 5 * 1024 * 1024
PARAGRAPH_INDENT = Pt(24)  # Alterado para pontos
MISTRAL_OCR_API_URL = "https://api.mistral.ai/v1/ocr"
MISTRAL_API_STATUS_URL = "https://api.mistral.ai/v1/status"

class SecurityException(Exception):    
    """Exceção para violações de políticas de segurança"""


def check_poppler_installed():
    """Verifica se o Poppler está instalado e disponível no PATH"""
    system = platform.system()
    
    try:
        if system == "Windows":
            # No Windows, verificamos se os binários do poppler estão no PATH
            paths = os.environ["PATH"].split(os.pathsep)
            poppler_found = any(
                os.path.exists(os.path.join(path, "pdftoppm.exe")) for path in paths
            )
            
            if not poppler_found:
                return False
        else:
            # No Linux/Mac, tentamos executar o comando
            with open(os.devnull, 'w') as devnull:
                subprocess.check_call(
                    ["pdftoppm", "-v"], 
                    stdout=devnull, 
                    stderr=devnull
                )
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False


def show_poppler_instructions():
    """Mostra instruções para instalação do Poppler baseado no sistema operacional"""
    system = platform.system()
    
    if system == "Windows":
        message = """
Poppler não encontrado! Siga as instruções para instalar:

1. Baixe o Poppler para Windows em: https://github.com/oschwartz10612/poppler-windows/releases/
2. Extraia os arquivos em uma pasta (ex: C:\\Poppler)
3. Adicione a pasta bin (ex: C:\\Poppler\\bin) ao PATH do sistema:
   a. Abra Painel de Controle > Sistema > Configurações avançadas do sistema
   b. Clique em "Variáveis de Ambiente"
   c. Em "Variáveis do Sistema", selecione "Path" e clique em "Editar"
   d. Adicione o caminho para a pasta bin
4. Reinicie o aplicativo

Alternativamente, instale o Poppler via Conda:
conda install -c conda-forge poppler
"""
    elif system == "Darwin":  # macOS
        message = """
Poppler não encontrado! Siga as instruções para instalar:

Para macOS, instale o Poppler via Homebrew:
brew install poppler

Após a instalação, reinicie o aplicativo.
"""
    else:  # Linux
        message = """
Poppler não encontrado! Siga as instruções para instalar:

Para Ubuntu/Debian:
sudo apt-get update
sudo apt-get install poppler-utils

Para Fedora/CentOS:
sudo dnf install poppler-utils

Após a instalação, reinicie o aplicativo.
"""
    
    if 'tkinter' in sys.modules:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror("Erro", message)
        root.destroy()
    else:
        print(message)


class BaseOCRProcessor:
    """Classe base para processadores de OCR"""
    
    def __init__(self):
        self.stop_event = threading.Event()
    
    @staticmethod
    def _validate_paths(input_dir: str, output_dir: str) -> None:
        """Validação de segurança dos caminhos"""
        if not os.path.exists(input_dir):
            raise SecurityException("O diretório de entrada não existe")
        
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir, exist_ok=True)
            except:
                raise SecurityException("Não foi possível criar o diretório de saída")
    
    def extract_text(self, pdf_path: str, lang: str = 'por') -> str:
        """Método a ser implementado pelas subclasses"""
        raise NotImplementedError("Método deve ser implementado pela subclasse")
    
    def get_paragraphs(self, text: str) -> List[Tuple[str, str]]:
        """
        Divide o texto em parágrafos preservando estrutura básica
        Retorna uma lista de tuplas (texto, tipo_paragrafo)
        """
        # Padrões para identificação de tipos de parágrafos
        artigo_pattern = re.compile(r'(artigo|art\.?)\s*\d+º?', re.IGNORECASE)
        titulo_pattern = re.compile(r'^(TÍTULO|CAPÍTULO|SEÇÃO)\s+[IVXLCDM0-9]+', re.IGNORECASE)
        
        # Dividir o texto em parágrafos potenciais
        raw_paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) 
                         if p.strip() and any(c.isalnum() for c in p) 
                         and len(p.strip()) > 10]
        
        processed_paragraphs = []
        
        for para in raw_paragraphs:
            # Substituir quebras de linha únicas por espaços para melhorar a leitura
            # Isso evita quebras de linha dentro de um mesmo parágrafo
            clean_para = re.sub(r'(?<!\n)\n(?!\n)', ' ', para).strip()
            
            # Identificar o tipo de parágrafo
            if titulo_pattern.search(clean_para):
                para_type = "titulo"
            elif artigo_pattern.search(clean_para):
                para_type = "artigo"
            elif "**" in clean_para:
                para_type = "destaque"
            else:
                para_type = "normal"
                
            processed_paragraphs.append((clean_para, para_type))
            
        return processed_paragraphs


class OCRProcessor(BaseOCRProcessor):
    """Classe responsável pelo processamento de OCR com Tesseract"""

    def __init__(self):
        super().__init__()
        self.poppler_available = check_poppler_installed()

    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Melhora a qualidade da imagem para OCR"""
        return ImageOps.autocontrast(
            image.convert('L').point(lambda x: 0 if x < 128 else 255))

    def extract_text(self, pdf_path: str, lang: str = 'por') -> str:
        """Extrai texto com fallback automático"""
        try:
            # Tenta extração direta primeiro
            text = extract_text(pdf_path)
            if len(text.strip()) > MIN_TEXT_LENGTH:
                return text

            # Verifica se o Poppler está disponível antes de tentar OCR
            if not self.poppler_available:
                logging.error("Poppler não está instalado ou não está no PATH")
                show_poppler_instructions()
                return text if len(text.strip()) > 0 else "Erro: Poppler não encontrado. Textos podem estar incompletos."

            # Fallback para OCR
            images = pdf2image.convert_from_path(pdf_path)
            return self._perform_ocr(images, lang)

        except Exception as e:
            logging.error(f"Erro no processamento: {e}")
            
            # Verificação específica para erro relacionado ao Poppler
            if "poppler" in str(e).lower():
                show_poppler_instructions()
                return "Erro: Poppler não encontrado ou não configurado corretamente."
            return ""

    def _perform_ocr(self, images: List[Image.Image], lang: str) -> str:
        """Executa OCR nas imagens do PDF"""
        text = ""
        for image in images:
            if self.stop_event.is_set():
                break

            processed = self._preprocess_image(image)
            hocr_data = pytesseract.image_to_pdf_or_hocr(
                processed,
                extension='hocr',
                config=f'--psm 1 -l {lang}'
            )
            
            soup = BeautifulSoup(hocr_data, 'html.parser')
            paragraphs = soup.find_all('p', class_='ocr_par')
            
            for para in paragraphs:
                lines = []
                for line in para.find_all('span', class_='ocr_line'):
                    words = line.find_all('span', class_='ocrx_word')
                    line_text = ' '.join(self._process_words(words))
                    lines.append(line_text)
                text += '\n'.join(lines) + '\n\n'

        return text if len(text.strip()) > MIN_TEXT_LENGTH else ""  

    def _process_words(self, words: List[BeautifulSoup]) -> List[str]:
        """Processa palavras extraídas do OCR e aplica formatações como negrito"""
        processed_words = []
        for word in words:
            word_text = word.get_text().strip()
            if 'bold' in word.get('class', []):  # Verifica se a palavra está em negrito
                processed_words.append(f"**{word_text}**")
            else:
                processed_words.append(word_text)
        return processed_words
        
MISTRAL_API_STATUS_URL = "https://api.mistral.ai/v1/status"  # Endpoint fictício para verificação
MISTRAL_endpoint_codestral_URL = "https://codestral.mistral.ai/v1/fim/completions" #endpoint real para codestral 
MISTRAL_endpoint_chat_URL = "https://codestral.mistral.ai/v1/chat/completions" #Endpoint de Chat  

class MistralOCRProcessor(BaseOCRProcessor):
    """Classe responsável pelo processamento de OCR com a API Mistral OCR"""
    
    def __init__(self, api_key=""):
        super().__init__()
        self.api_key = api_key
        self.api_url = MISTRAL_OCR_API_URL
        self.api_calls_count = 0
        self.total_tokens_used = 0
        self.active_requests = 0
        self.lock = threading.Lock()

    def extract_text(self, pdf_path: str, lang: str = 'por') -> str:
        """Extrai texto de um PDF usando a API Mistral OCR"""
        file_name = os.path.basename(pdf_path)
        log_prefix = f"[Mistral OCR][{file_name}]"
        
        try:
            if not self.api_key:
                logging.error(f"{log_prefix} API Key não configurada")
                return "Erro: API Key não configurada"
            
            if not os.path.exists(pdf_path):
                logging.error(f"{log_prefix} Arquivo não encontrado")
                return "Erro: Arquivo não encontrado"
            
            # Verificar se é um PDF válido
            with open(pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
                if not pdf_data.startswith(b'%PDF'):
                    logging.error(f"{log_prefix} O arquivo não é um PDF válido")
                    return "Erro: O arquivo não parece ser um PDF válido"
            
            # Tentar extração direta com pdfminer primeiro
            logging.info(f"{log_prefix} Tentando extração direta com pdfminer...")
            try:
                direct_text = extract_text(pdf_path)
                if len(direct_text.strip()) > MIN_TEXT_LENGTH:
                    logging.info(f"{log_prefix} Extração direta bem-sucedida: {len(direct_text)} caracteres")
                    return direct_text
                else:
                    logging.info(f"{log_prefix} Extração direta não forneceu texto suficiente. Usando Mistral OCR...")
            except Exception as e:
                logging.warning(f"{log_prefix} Erro na extração direta: {e}. Tentando Mistral OCR...")
                
            # Agora use a API Mistral OCR
            logging.info(f"{log_prefix} Iniciando processamento...")
            
            with open(pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
                return self._call_mistral_ocr_api(pdf_data, file_name, lang)
                
        except Exception as e:
            logging.error(f"{log_prefix} Erro: {str(e)}")
            return f"Erro: {str(e)}"

    def _call_mistral_ocr_api(self, pdf_data, file_name, lang: str) -> str:
        """Chama a API Mistral OCR para processar um PDF"""
        log_prefix = f"[Mistral OCR][{file_name}]"
        logging.info(f"{log_prefix} Iniciando processamento via API...")
        
        with self.lock:
            self.api_calls_count += 1
            self.active_requests += 1
        
        try:
            # Verificar se temos dados PDF válidos
            if not pdf_data or not pdf_data.startswith(b'%PDF'):
                logging.error(f"{log_prefix} Dados PDF inválidos.")
                return "Erro: Dados PDF inválidos."
                
            # Codificar o PDF como Base64 para envio
            base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
                
            # Tamanho do arquivo (para logs)
            file_size_kb = len(pdf_data) / 1024
            logging.info(f"{log_prefix} Tamanho do arquivo: {file_size_kb:.2f} KB")
            
            # Mapeamento de idioma para o formato esperado pela API
            lang_mapping = {
                'por': 'portuguese',
                'eng': 'english',
                'spa': 'spanish',
                'fra': 'french',
                'deu': 'german'
            }
            
            # Preparar payload para a API Mistral OCR
            payload = {
                "model": "mistral-ocr-latest",
                "id": str(uuid.uuid4()),
                "document": {
                    "type": "document_base64",
                    "document_base64": base64_pdf,
                    "document_name": file_name
                },
                "include_image_base64": False
            }
            
            # Adicionar idioma se disponível no mapeamento
            if lang in lang_mapping:
                payload["language"] = lang_mapping[lang]
            
            # Preparar headers com autenticação
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            }
            
            logging.info(f"{log_prefix} Enviando requisição para API Mistral OCR...")
            
            # Fazer a chamada à API
            response = requests.post(
                self.api_url,
                headers=headers,
                json=payload,
                timeout=120  # Timeout aumentado para arquivos grandes
            )
            
            # Verificar resposta
            # Verificar resposta
            if response.status_code == 200:
                logging.info(f"{log_prefix} Requisição bem-sucedida. Status: 200")
                
                # Processar resposta
                result = response.json()
                
                # Extrai e concatena texto de todas as páginas
                all_text = ""
                
                # Extrair texto de acordo com a estrutura da resposta da API
                if "pages" in result:
                    pages = result["pages"]
                    pages_count = len(pages)
                    logging.info(f"{log_prefix} Processadas {pages_count} páginas")
                    
                    for i, page in enumerate(pages, 1):
                        if "text" in page:
                            page_text = page["text"].strip()
                            all_text += page_text + "\n\n"
                            logging.debug(f"{log_prefix} Página {i}: {len(page_text)} caracteres")
                        
                        # Ou se a API retorna "markdown" em vez de "text"
                        elif "markdown" in page:
                            page_text = page["markdown"].strip()
                            all_text += page_text + "\n\n"
                    
                    # Atualizar contagem de tokens
                    if "usage_info" in result:
                        with self.lock:
                            if "pages_processed" in result["usage_info"]:
                                self.total_tokens_used += result["usage_info"]["pages_processed"]
                            # Forçar atualização da interface na thread principal
                            if hasattr(self.app, '_update_api_stats'):
                                self.app._update_api_stats()
                
                # Verificar se temos texto significativo
                if all_text and len(all_text.strip()) > MIN_TEXT_LENGTH:
                    return all_text.strip()
                else:
                    logging.warning(f"{log_prefix} Texto extraído muito curto ou vazio.")
                    return "Nenhum texto significativo extraído pelo Mistral OCR."
            else:
                error_message = f"Erro na API Mistral OCR: Status {response.status_code}"
                try:
                    error_details = response.json()
                    error_message += f" - {error_details.get('error', {}).get('message', '')}"
                except:
                    pass
                    
                logging.error(f"{log_prefix} {error_message}")
                return error_message
                
        except Exception as e:
            logging.error(f"{log_prefix} Exceção: {str(e)}")
            return f"Erro ao processar com Mistral OCR: {str(e)}"
            
        finally:
            with self.lock:
                self.active_requests -= 1

class PDFProcessorApp(tk.Tk):
    def _create_api_status_frame(self):
        """Nova seção para monitoramento da API"""
        status_frame = ttk.LabelFrame(self, text="Status da API")
        status_frame.pack(fill='x', padx=10, pady=5)

        # Grid para organização
        status_frame.grid_columnconfigure(1, weight=1)

        # Labels dinâmicos
        ttk.Label(status_frame, text="Conexão:").grid(row=0, column=0, sticky='w')
        self.api_connection_label = ttk.Label(status_frame, text="Desconectado", foreground="red")
        self.api_connection_label.grid(row=0, column=1, sticky='w')

        ttk.Label(status_frame, text="Requisições Ativas:").grid(row=1, column=0, sticky='w')
        self.active_requests_label = ttk.Label(status_frame, text="0")
        self.active_requests_label.grid(row=1, column=1, sticky='w')

        ttk.Label(status_frame, text="Tokens Usados:").grid(row=2, column=0, sticky='w')
        self.tokens_used_label = ttk.Label(status_frame, text="0")
        self.tokens_used_label.grid(row=2, column=1, sticky='w')

        ttk.Label(status_frame, text="Chamadas Totais:").grid(row=3, column=0, sticky='w')
        self.total_calls_label = ttk.Label(status_frame, text="0")
        self.total_calls_label.grid(row=3, column=1, sticky='w')

        # Botão para atualizar status
        ttk.Button(status_frame, text="Atualizar",).grid(row=4, columnspan=2)
        # Botão para atualizar status - precisa de um comando
        ttk.Button(status_frame, text="Atualizar", command=self._update_api_stats).grid(row=4, columnspan=2)

    def _test_mistral_api(self):
        """Testa a conexão com a API Mistral OCR usando apenas verificação de autenticação"""
        api_key = self.api_key_entry.get().strip()
        
        if not api_key:
            messagebox.showwarning("Aviso", "Insira uma API Key para testar a conexão!")
            return
        
        # Atualiza status para indicar teste em andamento
        self.api_status_label.config(
            text="Status: Testando conexão...",
            foreground="blue"
        )
        self.update_idletasks()
        
        try:
            # Headers com autenticação
            headers = {
                "Authorization": f"Bearer {api_key}"
            }
            
            # Usamos apenas o endpoint de modelos para verificar a autenticação
            response = requests.get(
                "https://api.mistral.ai/v1/models",
                headers=headers,
                timeout=10
            )
            
            # Verificar resposta
            if response.status_code in [200, 201, 204] or response.status_code == 422:
                self.api_status_label.config(
                    text="Status: Conexão bem-sucedida! API pronta para uso.",
                    foreground="green"
                )
                messagebox.showinfo("Sucesso", "Conexão com a API Mistral estabelecida com sucesso!")
                
                # Atualiza variáveis de estado
                self.mistral_ocr.api_key = api_key
                self.api_connection_label.config(text="Conectado", foreground="green")
                return True
            
                # Atualizar estado de conexão
                self.api_connection_label.config(text="Conectado", foreground="green")

                # Forçar atualização das estatísticas
                self._update_api_stats()
                
                return True
                
            elif response.status_code == 401:
                self.api_status_label.config(
                    text="Status: Falha de autenticação - API Key inválida!",
                    foreground="red"
                )
                messagebox.showerror("Erro", "API Key inválida ou expirada. Verifique suas credenciais.")
                return False
                
            else:
                self.api_status_label.config(
                    text=f"Status: Erro na conexão - Código {response.status_code}",
                    foreground="red"
                )
                messagebox.showerror("Erro", f"Erro ao conectar com a API Mistral. Código: {response.status_code}")
                return False
                
        except Exception as e:
            self.api_status_label.config(
                text=f"Status: Erro desconhecido na conexão",
                foreground="red"
            )
            messagebox.showerror("Erro", f"Erro ao testar conexão: {str(e)}")
            return False
            
    def extract_text(self, pdf_path: str, lang: str = 'por') -> str:
        """Extrai texto de um PDF usando a API Mistral OCR"""
        file_name = os.path.basename(pdf_path)
        log_prefix = f"[Mistral OCR][{file_name}]"
        
        try:
            if not self.api_key:
                logging.error(f"{log_prefix} API Key não configurada")
                return "Erro: API Key não configurada"
            
            if not os.path.exists(pdf_path):
                logging.error(f"{log_prefix} Arquivo não encontrado")
                return "Erro: Arquivo não encontrado"
            
            # Verificar se é um PDF válido
            with open(pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
                if not pdf_data.startswith(b'%PDF'):
                    logging.error(f"{log_prefix} O arquivo não é um PDF válido")
                    return "Erro: O arquivo não parece ser um PDF válido"
            
            # Tentar extração direta com pdfminer primeiro
            logging.info(f"{log_prefix} Tentando extração direta com pdfminer...")
            try:
                direct_text = extract_text(pdf_path)
                if len(direct_text.strip()) > MIN_TEXT_LENGTH:
                    logging.info(f"{log_prefix} Extração direta bem-sucedida: {len(direct_text)} caracteres")
                    return direct_text
                else:
                    logging.info(f"{log_prefix} Extração direta não forneceu texto suficiente. Usando Mistral OCR...")
            except Exception as e:
                logging.warning(f"{log_prefix} Erro na extração direta: {e}. Tentando Mistral OCR...")
                
            # Agora use a API Mistral OCR
            logging.info(f"{log_prefix} Iniciando processamento...")
            
            with open(pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
                return self._call_mistral_ocr_api(pdf_data, file_name, lang)
             
        except Exception as e:
            logging.error(f"{log_prefix} Erro: {str(e)}")
            return f"Erro: {str(e)}"
    
    def _load_api_key_from_file(self) -> Optional[str]:
        """Carrega a chave da API do arquivo"""
        try:
            # Diretório de configuração
            config_dir = os.path.join(os.path.expanduser("~"), ".ocr_processor")
            key_file = os.path.join(config_dir, "api_key.dat")
            
            if not os.path.exists(key_file):
                logging.info("Arquivo de chave da API não encontrado")
                return None
            
            # Descriptografar a chave
            import base64
            
            with open(key_file, 'rb') as f:
                encoded_data = f.read()
                
            # Decodificar
            decoded_data = base64.b64decode(encoded_data)
            
            # Remover o salt (16 bytes iniciais) e converter para string
            api_key = decoded_data[16:].decode('utf-8')
            
            logging.info("Chave da API carregada com sucesso")
            return api_key
        
        except Exception as e:
            logging.error(f"Erro ao carregar chave da API: {e}")
            return None
    
class DocxFormatter:
    """Classe para formatação avançada de documentos DOCX"""
    
    @staticmethod
    def sanitize_text_for_xml(text: str) -> str:
        """
        Limpa o texto de caracteres incompatíveis com XML
        
        Args:
            text: O texto a ser limpo
            
        Returns:
            str: Texto limpo compatível com XML
        """
        if not text:
            return ""
            
        # Função para verificar se um caractere é válido para XML
        def is_xml_char(c):
            # XML aceita: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
            cp = ord(c)
            return (
                cp == 0x9 or
                cp == 0xA or
                cp == 0xD or
                (0x20 <= cp <= 0xD7FF) or
                (0xE000 <= cp <= 0xFFFD) or
                (0x10000 <= cp <= 0x10FFFF)
            )
        
        # Filtrar apenas caracteres válidos para XML
        return ''.join(c for c in text if is_xml_char(c))
    
    @staticmethod
    def setup_document_styles(doc: Document) -> None:
        """Configura estilos do documento para melhor formatação de parágrafos"""
        # Estilo para parágrafos normais
        if 'Normal Paragraph' not in doc.styles:
            normal_style = doc.styles.add_style('Normal Paragraph', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            normal_style.paragraph_format.first_line_indent = Inches(PARAGRAPH_INDENT)
            normal_style.paragraph_format.space_after = Pt(10)
        
        # Estilo para artigos
        if 'Artigo' not in doc.styles:
            artigo_style = doc.styles.add_style('Artigo', WD_STYLE_TYPE.PARAGRAPH)
            artigo_style.font.size = Pt(12)
            artigo_style.font.bold = True
            artigo_style.paragraph_format.space_before = Pt(12)
            artigo_style.paragraph_format.space_after = Pt(6)
        
        # Estilo para texto destacado
        if 'Destaque' not in doc.styles:
            destaque_style = doc.styles.add_style('Destaque', WD_STYLE_TYPE.PARAGRAPH)
            destaque_style.font.size = Pt(12)
            destaque_style.font.bold = True
            destaque_style.paragraph_format.first_line_indent = Inches(PARAGRAPH_INDENT)
            destaque_style.paragraph_format.space_after = Pt(10)
    
    @staticmethod
    def add_paragraph_with_style(doc: Document, text: str, para_type: str) -> None:
        """Adiciona parágrafo com o estilo apropriado baseado no tipo de parágrafo"""
        text = unescape(text)  # Decodifica entidades HTML
        
        # Remove marcadores de negrito se presentes
        clean_text = text.replace('**', '')
        
        # Sanitiza o texto para garantir compatibilidade com XML
        clean_text = DocxFormatter.sanitize_text_for_xml(clean_text)
        
        if not clean_text:  # Se após a limpeza o texto ficou vazio, ignore
            return
        
        try:
            if para_type == "titulo":
                doc.add_heading(clean_text, level=1)
            elif para_type == "artigo":
                p = doc.add_paragraph(clean_text, style='Artigo')
            elif para_type == "destaque":
                p = doc.add_paragraph(clean_text, style='Destaque')
            else:  # normal
                p = doc.add_paragraph(clean_text, style='Normal Paragraph')
        except ValueError as e:
            # Se ainda houver erro, log o problema e pula este parágrafo
            logging.warning(f"Não foi possível adicionar um parágrafo: {e}. Texto: {clean_text[:50]}...")


class JsonFormatter:
    @staticmethod
    def create_mistral_entry(text: str, paragraphs: List[Tuple[str, str]]) -> Optional[Dict]:
        """Cria entrada garantindo que termina com 'assistant'"""
        messages = []
        
        # Primeira mensagem sempre do usuário
        messages.append({
            "role": "user",
            "content": JsonFormatter.sanitize_text(text[:5000])  # Limite de contexto
        })
        
        # Resposta do assistente com parágrafos relevantes
        assistant_content = "\n\n".join([p[0] for p in paragraphs if p[0].strip()])
        
        if not assistant_content:  # Ignora entradas sem resposta
            return None
            
        messages.append({
            "role": "assistant",
            "content": JsonFormatter.sanitize_text(assistant_content)
        })
        
        return {"messages": messages} if len(messages) >= 2 else None



    @staticmethod
    def sanitize_text(text: str) -> str:
        """Sanitização mais rigorosa para compatibilidade com LLMs"""
        # Remove caracteres especiais e normaliza espaços
        cleaned = re.sub(r'\s+', ' ', text.strip())
        # Remove caracteres não-ASCII
        return cleaned.encode('ascii', 'ignore').decode()

class PDFProcessorApp(tk.Tk):
    """Interface gráfica principal"""

    def __init__(self):
        super().__init__()
        self.title("PDF Processor Pro v5 - Tesseract & Mistral OCR")
        self.geometry("800x700")
        
        # Inicializar processadores OCR
        self.tesseract_ocr = OCRProcessor()
        self.mistral_ocr = MistralOCRProcessor()
        self.current_ocr = self.tesseract_ocr  # Default para Tesseract
        
         # Inicializar tratador de erros
        self.error_handler = DocumentErrorHandler()
        
        # Adicionar o lock para operações de escrita JSON
        self._json_write_lock = threading.Lock()
        
        # Verifica o Poppler para Tesseract
        if not self.tesseract_ocr.poppler_available:
            show_poppler_instructions()
                
        self._setup_ui()
        # Iniciar atualização periódica das estatísticas da API
        self._update_api_stats()

        # Adicionar opção para geração de sumário
        self.generate_summary_var = tk.BooleanVar(value=False)
        summary_frame = ttk.Frame(self)
        summary_frame.pack(fill='x', padx=10, pady=5)
        ttk.Checkbutton(
            summary_frame, 
            text="Gerar sumário e tabela de conteúdo usando IA", 
            variable=self.generate_summary_var
        ).pack(anchor='w')

    def _setup_ui(self):
        """Configura componentes da interface com log na lateral direita"""
        self.input_dir_var = tk.StringVar()
        self.output_dir_var = tk.StringVar()
        self.progress_var = tk.DoubleVar()
        self.lang_var = tk.StringVar(value='por')
        self.ocr_type_var = tk.StringVar(value='tesseract')  # Default para Tesseract
        
        # Frame principal para dividir a interface em dois lados
        main_frame = ttk.Frame(self)
        main_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Frame para os controles (lado esquerdo)
        left_frame = ttk.Frame(main_frame)
        left_frame.pack(side='left', fill='both', expand=True)
        
        # Frame para log (lado direito)
        right_frame = ttk.LabelFrame(main_frame, text="Log de Operações")
        right_frame.pack(side='right', fill='both', expand=True, padx=5, pady=5)
        
        # Configuração do log no frame da direita
        self.log_text = scrolledtext.ScrolledText(right_frame, state='disabled', width=40, height=30)
        self.log_text.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Configuração de logging
        logging.basicConfig(level=logging.INFO,
                        handlers=[self._create_file_handler(),
                                self._create_gui_handler(self.log_text)])
        
        # Adiciona os componentes ao frame da esquerda
        self._create_api_status_frame(left_frame)
        self._create_directory_selector(left_frame)
        self._create_ocr_type_selector(left_frame)
        self._create_language_selector(left_frame)
        
        # Adicionar opção para geração de sumário
        self.generate_summary_var = tk.BooleanVar(value=False)
        summary_frame = ttk.Frame(left_frame)
        summary_frame.pack(fill='x', padx=10, pady=5)
        ttk.Checkbutton(
            summary_frame, 
            text="Gerar sumário e tabela de conteúdo usando IA", 
            variable=self.generate_summary_var
        ).pack(anchor='w')
        
        # Adicionar opção para extração de dados estruturados
        self.extract_data_var = tk.BooleanVar(value=False)
        extract_frame = ttk.Frame(left_frame)
        extract_frame.pack(fill='x', padx=10, pady=5)
        ttk.Checkbutton(
            extract_frame, 
            text="Extrair dados estruturados dos documentos", 
            variable=self.extract_data_var
        ).pack(anchor='w')
        
        self._create_controls(left_frame)
        self._create_progress_bar(left_frame)

        
    def _create_api_status_frame(self, parent_frame):
        """Nova seção para monitoramento da API"""
        status_frame = ttk.LabelFrame(parent_frame, text="Status da API")
        status_frame.pack(fill='x', padx=10, pady=5)
        
        # Grid para organização
        status_frame.grid_columnconfigure(1, weight=1)
        
        # Labels dinâmicos
        ttk.Label(status_frame, text="Conexão:").grid(row=0, column=0, sticky='w')
        self.api_connection_label = ttk.Label(status_frame, text="Desconectado", foreground="red")
        self.api_connection_label.grid(row=0, column=1, sticky='w')

    def _create_directory_selector(self, parent_frame):
        """Componentes de seleção de diretórios"""
        dir_frame = ttk.Frame(parent_frame)
        dir_frame.pack(fill='x', padx=10, pady=5)

        for label, var in [("Entrada:", self.input_dir_var),
                        ("Saída:", self.output_dir_var)]:
            frame = ttk.Frame(dir_frame)
            frame.pack(fill='x', pady=2)

            ttk.Label(frame, text=label).pack(side='left')
            ttk.Entry(frame, textvariable=var, width=40).pack(side='left', expand=True)
            ttk.Button(frame, text="📁", command=lambda v=var: v.set(filedialog.askdirectory()))\
                .pack(side='left')

    def _create_ocr_type_selector(self, parent_frame):
        """Seletor de tipo de OCR com opção para API em arquivo"""
        ocr_frame = ttk.LabelFrame(parent_frame, text="Tipo de OCR")
        ocr_frame.pack(fill='x', padx=10, pady=5)
        
        # Opção Tesseract
        ttk.Radiobutton(
            ocr_frame, 
            text="Tesseract OCR (local)", 
            variable=self.ocr_type_var,
            value="tesseract",
            command=self._update_ocr_processor
        ).pack(anchor='w', padx=10)
        
        # Opção Mistral com API digitada
        ttk.Radiobutton(
            ocr_frame, 
            text="Mistral OCR (digitar API Key)", 
            variable=self.ocr_type_var,
            value="mistral",
            command=self._update_ocr_processor
        ).pack(anchor='w', padx=10)
        
        # Nova opção: Mistral com API de arquivo
        ttk.Radiobutton(
            ocr_frame, 
            text="Mistral OCR (usar API Key em arquivo)", 
            variable=self.ocr_type_var,
            value="mistral_file",
            command=self._update_ocr_processor
        ).pack(anchor='w', padx=10)
        
        # Frame para configurações do Mistral
        self.mistral_config_frame = ttk.Frame(ocr_frame)
        self.mistral_config_frame.pack(fill='x', pady=5)
        
        # [resto do código do método]
        
    def _create_language_selector(self, parent_frame):
        """Seletor de idioma para OCR"""
        lang_frame = ttk.Frame(parent_frame)
        lang_frame.pack(pady=5)

        ttk.Label(lang_frame, text="Idioma:").pack(side='left')
        ttk.Combobox(lang_frame, textvariable=self.lang_var,
                values=SUPPORTED_LANGS, state='readonly').pack(side='left')

    def _create_controls(self, parent_frame):
        """Botões de controle"""
        btn_frame = ttk.Frame(parent_frame)
        btn_frame.pack(pady=10)

        ttk.Button(btn_frame, text="Iniciar", command=self._start_processing).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="Cancelar", command=self._cancel_processing).pack(side='left', padx=5)
        
    def _create_progress_bar(self, parent_frame):
        """Barra de progresso e status"""
        progress_frame = ttk.Frame(parent_frame)
        progress_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Label(progress_frame, text="Progresso:").pack(side='left')
        
        # Barra de progresso
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100, length=300)
        self.progress_bar.pack(side='left', fill='x', expand=True, padx=5)
        
        # Texto de porcentagem
        self.progress_label = ttk.Label(progress_frame, textvariable=self.progress_var)
        self.progress_label.pack(side='left')
        ttk.Label(progress_frame, text="%").pack(side='left')
        
    
    def _update_api_key(self):
        """Atualiza a API key do Mistral OCR"""
        new_key = self.api_key_entry.get().strip()
        if new_key:
            self.mistral_ocr.api_key = new_key
            self.api_status_label.config(
                text="Status: API key configurada (não testada)",
                foreground="blue"
            )
            messagebox.showinfo("Sucesso", "API Key atualizada com sucesso! Recomendamos testar a conexão.")
        else:
            messagebox.showwarning("Aviso", "API Key não pode estar vazia.")
            self.api_status_label.config(
                text="Status: API key não configurada",
                foreground="orange"
            )
    
    def _test_mistral_api(self):
        """Testa a conexão com a API Mistral OCR"""
        api_key = self.api_key_entry.get().strip()
        
        if not api_key:
            messagebox.showwarning("Aviso", "Insira uma API Key para testar a conexão!")
            return
        
        # Atualiza status para indicar teste em andamento
        self.api_status_label.config(
            text="Status: Testando conexão...",
            foreground="blue"
        )
        self.update_idletasks()  # Atualiza a UI imediatamente
        
        try:
            # Preparar headers com autenticação
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            # Criar um payload válido para teste (requerido pela API Mistral)
            test_payload = {
                "model": "mistral-ocr-latest",
                "id": str(uuid.uuid4()),
                "document": {
                    "type": "document_base64",
                    "document_base64": "SGVsbG8gV29ybGQ=",  # "Hello World" em base64
                    "document_name": "test.txt"
                }
            }
            
            # Fazer uma requisição POST correta
            response = requests.post(
                MISTRAL_OCR_API_URL,
                headers=headers,
                json=test_payload,  # Usando json parameter para converter automaticamente para JSON
                timeout=10
            )
            
            # Verificar resposta
            if response.status_code == 200:
                self.api_status_label.config(
                    text="Status: Conexão bem-sucedida! API pronta para uso.",
                    foreground="green"
                )
                messagebox.showinfo("Sucesso", "Conexão com a API Mistral OCR estabelecida com sucesso!")
            elif response.status_code == 401:
                self.api_status_label.config(
                    text="Status: Falha de autenticação - API Key inválida!",
                    foreground="red"
                )
                messagebox.showerror("Erro", "API Key inválida ou expirada. Verifique suas credenciais.")
            elif response.status_code == 422:
                # Este status pode indicar erro de validação mas significa que a autenticação funcionou
                self.api_status_label.config(
                    text="Status: API Key válida, mas requisição de teste precisa de ajustes.",
                    foreground="green"
                )
                messagebox.showinfo("Parcialmente Sucesso", "API Key aceita, mas a estrutura da requisição precisa de ajustes.")
            else:
                self.api_status_label.config(
                    text=f"Status: Erro na conexão - Código {response.status_code}",
                    foreground="red"
                )
                messagebox.showerror("Erro", f"Erro ao conectar com a API Mistral OCR. Código: {response.status_code}")
                
        except requests.exceptions.ConnectionError:
            self.api_status_label.config(
                text="Status: Falha de conexão - Verifique sua internet",
                foreground="red"
            )
            messagebox.showerror("Erro", "Não foi possível conectar ao servidor da API. Verifique sua conexão de internet.")
        except requests.exceptions.Timeout:
            self.api_status_label.config(
                text="Status: Timeout na conexão com a API",
                foreground="red"
            )
            messagebox.showerror("Erro", "Tempo de conexão esgotado. O servidor pode estar sobrecarregado.")
        except Exception as e:
            self.api_status_label.config(
                text=f"Status: Erro desconhecido na conexão",
                foreground="red"
            )
            messagebox.showerror("Erro", f"Erro ao testar conexão: {str(e)}")
    
    def _update_ocr_processor(self):
        """Atualiza o processador OCR baseado na seleção do usuário"""
        ocr_type = self.ocr_type_var.get()
        
        if ocr_type == "tesseract":
            self.current_ocr = self.tesseract_ocr
            logging.info("Usando Tesseract OCR para processamento")
            # Ocultar configurações específicas do Mistral
            self.mistral_config_frame.pack_forget()
            
        elif ocr_type == "mistral":
            # Atualizar a API key antes de usar
            self.mistral_ocr.api_key = self.api_key_entry.get().strip()
            self.current_ocr = self.mistral_ocr
            logging.info("Usando Mistral OCR para processamento")
            # Mostrar configurações específicas do Mistral
            self.mistral_config_frame.pack(fill='x', pady=5)
            
        elif ocr_type == "mistral_file":
            # Carregar a chave da API do arquivo
            api_key = self._load_api_key_from_file()
            if api_key:
                self.mistral_ocr.api_key = api_key
                # Atualizar o campo de entrada para mostrar que uma chave foi carregada 
                # (sem revelar a chave completa)
                self.api_key_entry.delete(0, tk.END)
                self.api_key_entry.insert(0, "********" + api_key[-4:])
                self.current_ocr = self.mistral_ocr
                logging.info("Usando Mistral OCR com chave carregada do arquivo")
                # Mostrar configurações específicas do Mistral
                self.mistral_config_frame.pack(fill='x', pady=5)
                # Atualizar status
                self.api_status_label.config(
                    text="Status: API key carregada do arquivo",
                    foreground="green"
                )
            else:
                messagebox.showwarning("Aviso", "Não foi possível carregar a chave da API do arquivo. "
                                    "Por favor, digite uma chave manualmente.")
                # Mudar para modo manual
                self.ocr_type_var.set("mistral")
                self.mistral_config_frame.pack(fill='x', pady=5)
                self.api_status_label.config(
                    text="Status: Falha ao carregar chave do arquivo",
                    foreground="red"
                )
        
        # Atualizar status da conexão
        if ocr_type in ["mistral", "mistral_file"] and self.mistral_ocr.api_key:
            self.api_status_label.config(
                text="Status: API key configurada, pronta para processamento",
                foreground="green"
            )
        elif ocr_type in ["mistral", "mistral_file"]:
            self.api_status_label.config(
                text="Status: API key não configurada",
                foreground="orange"
            )
      
    def _cancel_processing(self):
        """Cancela o processamento em andamento"""
        self.current_ocr.stop_event.set()
        logging.info("Processamento cancelado pelo usuário")

    def _create_file_handler(self):
        """Handler para arquivo de log"""
        return RotatingFileHandler('processing_ref_passo1ocrcommistral0adequa-dataset_mistral.log', maxBytes=MAX_LOG_SIZE,
                                 backupCount=5, encoding='utf-8')

    def _create_gui_handler(self, widget):
        """Handler para exibição na interface"""
        class GuiHandler(logging.Handler):
            def emit(self, record):
                widget.configure(state='normal')
                widget.insert(tk.END, self.format(record) + '\n')
                widget.see(tk.END)  # Auto-scroll para o fim
                widget.configure(state='disabled')
        
        handler = GuiHandler()
        handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
        return handler

    def _start_processing(self):
        """Inicia o processamento em thread separada"""
        input_dir = self.input_dir_var.get()
        output_dir = self.output_dir_var.get()

        if not (input_dir and output_dir):
            messagebox.showwarning("Aviso", "Selecione os diretórios de entrada e saída!")
            return
        
        ocr_type = self.ocr_type_var.get()
        
        # Verificações específicas para cada modo OCR
        if ocr_type == "mistral":
            # Verificar API key se usando Mistral
            if not self.api_key_entry.get().strip():
                messagebox.showwarning("Aviso", "API Key do Mistral OCR não configurada. Por favor, configure uma chave válida.")
                return        
        elif ocr_type == "tesseract":
            # Verificar o Poppler se usando Tesseract
            if not self.tesseract_ocr.poppler_available:
                if messagebox.askyesno("Aviso", 
                                    "O Poppler não está instalado, o que pode causar problemas na extração de texto de PDFs. Deseja continuar mesmo assim?"):
                    pass
                else:
                    show_poppler_instructions()
                    return

        # Resetar eventos de parada
        self.tesseract_ocr.stop_event.clear()
        self.mistral_ocr.stop_event.clear()
        
        # Atualizar processador
        self._update_ocr_processor()
        
        # Atualizar UI
        self.progress_var.set(0)
        logging.info(f"Iniciando processamento com {ocr_type.upper()} OCR")
        
        # Iniciar processamento em thread separada
        processing_thread = threading.Thread(
            target=self._process_files,
            args=(input_dir, output_dir),
            daemon=True
        )
        processing_thread.start()

    def _process_files(self, input_dir: str, output_dir: str):
        """Processa todos os arquivos PDF e imagens"""
        try:
            self.current_ocr._validate_paths(input_dir, output_dir)
            
            # Lista todos os PDFs e imagens suportadas
            pdfs = [f for f in os.listdir(input_dir) if f.lower().endswith('.pdf')]
            images = [f for f in os.listdir(input_dir) if f.lower().endswith(('.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp'))]
            
            all_files = pdfs + images
            
            if not all_files:
                messagebox.showinfo("Informação", "Nenhum arquivo PDF ou imagem encontrado no diretório de entrada.")
                return
            
            logging.info(f"Iniciando processamento de {len(all_files)} arquivos ({len(pdfs)} PDFs, {len(images)} imagens) com {self.ocr_type_var.get()} OCR")

            # Resultados de processamento individual
            processed_files = []
            extracted_data = []

            with ThreadPoolExecutor() as executor:
                futures = {executor.submit(self._process_single_file_or_image,
                                        os.path.join(input_dir, f),
                                        output_dir): f for f in all_files}

                for i, future in enumerate(as_completed(futures), 1):
                    file_name = futures[future]
                    result = future.result()
                    processed_files.append((os.path.join(input_dir, file_name), result))
                    
                    # Atualiza o progresso
                    progress = (i / len(all_files)) * 100
                    self.progress_var.set(round(progress, 1))
                    self.update_idletasks()  # Atualiza a interface
                    
                    if self.current_ocr.stop_event.is_set():
                        break

            # Extração de dados estruturados (se habilitada)
            if self.extract_data_var.get() and hasattr(self, 'mistral_ocr') and self.mistral_ocr.api_key:
                logging.info("Iniciando extração de dados estruturados...")
                
                # Lista de arquivos processados com sucesso
                successful_files = [path for path, success in processed_files if success]
                
                if successful_files:
                    try:
                        # Cria o extrator
                        extractor = DocumentDataExtractor(self.mistral_ocr.api_key)
                        
                        # Processa os documentos
                        data_df = extractor.process_document_batch(successful_files)
                        
                        # Salva CSV com dados extraídos
                        csv_path = os.path.join(output_dir, "dados_extraidos.csv")
                        data_df.to_csv(csv_path, index=False)
                        
                        logging.info(f"Dados estruturados salvos em {csv_path}")
                        messagebox.showinfo("Extração de Dados", f"Dados estruturados extraídos de {len(successful_files)} documentos e salvos em {csv_path}")
                    except Exception as e:
                        logging.error(f"Erro na extração de dados: {e}")
                        messagebox.showerror("Erro", f"Falha na extração de dados estruturados: {e}")
                else:
                    logging.warning("Nenhum arquivo processado com sucesso para extração de dados.")

            # Mensagem de conclusão
            if not self.current_ocr.stop_event.is_set():
                messagebox.showinfo("Concluído", f"Processamento concluído com sucesso! {i} de {len(all_files)} arquivos processados.")
            else:
                messagebox.showinfo("Interrompido", f"Operação interrompida. {i} de {len(all_files)} arquivos processados.")

        except SecurityException as se:
            messagebox.showerror("Erro de Segurança", str(se))
        except Exception as e:
            logging.error(f"Erro crítico: {e}")
            messagebox.showerror("Erro", f"Falha no processamento: {e}")

    def _process_single_file_or_image(self, file_path: str, output_dir: str) -> bool:
        """
        Processa um único arquivo, seja PDF ou imagem
        
        Args:
            file_path (str): Caminho do arquivo a ser processado
            output_dir (str): Diretório de saída
        
        Returns:
            bool: True se processado com sucesso, False caso contrário
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Verificar se é uma imagem
        if file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp']:
            return self._process_single_image(file_path, output_dir)
        else:  # PDF ou outro formato de documento
            return self._process_single_file(file_path, output_dir)
    
    def _process_single_image(self, image_path: str, output_dir: str) -> bool:
        """
        Processa uma única imagem
        
        Args:
            image_path (str): Caminho da imagem
            output_dir (str): Diretório de saída
        
        Returns:
            bool: True se processado com sucesso, False caso contrário
        """
        try:
            file_name = os.path.basename(image_path)
            logging.info(f"Processando imagem {file_name}")
            
            # Abrir a imagem
            with Image.open(image_path) as img:
                # Pré-processar a imagem
                preprocessed = self.tesseract_ocr._preprocess_image(img)
                
                # Extrair texto usando Tesseract (mesmo para Mistral, usamos Tesseract para imagens)
                text = pytesseract.image_to_string(preprocessed, lang=self.lang_var.get())
                
                if not text or len(text.strip()) < MIN_TEXT_LENGTH:
                    logging.warning(f"Texto extraído da imagem {file_name} é muito curto ou vazio.")
                    return False
                
                # Processar parágrafos
                paragraphs = self.current_ocr.get_paragraphs(text)
                
                # Gerar DOCX
                docx_path = os.path.join(output_dir, f"{os.path.splitext(file_name)[0]}.docx")
                
                # Criar documento
                doc = Document()
                
                # Configurar estilos
                DocxFormatter.setup_document_styles(doc)
                
                # Adicionar título
                doc.add_heading(f"Imagem: {file_name}", level=0)
                
                # Adicionar metadados
                doc.add_paragraph(f"Processado com: Tesseract OCR")
                doc.add_paragraph(f"Data de processamento: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
                doc.add_paragraph(f"Idioma: {self.lang_var.get()}")
                doc.add_paragraph("").paragraph_format.space_after = Pt(20)
                
                # Adicionar parágrafos
                success_count = 0
                for para_text, para_type in paragraphs:
                    try:
                        DocxFormatter.add_paragraph_with_style(doc, para_text, para_type)
                        success_count += 1
                    except Exception as e:
                        logging.warning(f"Erro ao processar parágrafo da imagem: {str(e)[:100]}...")
                
                # Salvar documento
                doc.save(docx_path)
                logging.info(f"Documento DOCX criado para imagem: {docx_path}")
                
                # Gerar JSON se necessário
                json_success = self._generate_json(image_path, output_dir, text, paragraphs)
                
                # Gerar sumário e tabela de conteúdo, se solicitado
                if self.generate_summary_var.get():
                    self._generate_summary_and_toc(docx_path, text)
                
                return True
                
        except Exception as e:
            logging.error(f"Erro ao processar imagem {image_path}: {e}")
            error_type = self.error_handler.handle_extraction_error(e, {"file": image_path})
            recovery_action = self.error_handler.suggest_recovery_action(error_type, image_path)
            logging.info(f"Sugestão de recuperação: {recovery_action['action']}")
            return False
    
    def _process_single_file(self, file_path: str, output_dir: str):
        """Processa um único arquivo PDF com tratamento robusto de erros"""
        try:
            file_name = os.path.basename(file_path)
            logging.info(f"Processando {file_name} com {self.ocr_type_var.get()} OCR")
            
            # Extrair texto usando o OCR selecionado
            text = self.current_ocr.extract_text(file_path, self.lang_var.get())
            
            # Verificar se houve erro ou se o texto está vazio
            if not text or text.startswith("Erro:"):
                logging.error(f"Falha ao extrair texto de {file_name}: {text}")
                return False
            
            # Processar parágrafos com metadados
            paragraphs = self.current_ocr.get_paragraphs(text)
            
            # Geração do arquivo DOCX com parágrafos formatados
            docx_success = self._generate_docx(file_path, output_dir, paragraphs)
            
            # Geração do arquivo JSON estruturado
            json_success = self._generate_json(file_path, output_dir, text, paragraphs)
            
            # Gerar sumário e tabela de conteúdo, se solicitado
            if self.generate_summary_var.get() and docx_success:
                docx_path = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(file_path))[0]}.docx")
                self._generate_summary_and_toc(docx_path, text)
            
            return docx_success and json_success
            
        except UnicodeDecodeError as ude:
            logging.error(f"Erro de codificação ao processar {file_path}: {ude}")
            error_type = self.error_handler.handle_extraction_error(ude, {"file": file_path})
            recovery_action = self.error_handler.suggest_recovery_action(error_type, file_path)
            logging.info(f"Sugestão de recuperação: {recovery_action['action']}")
            return False
            
        except PermissionError as pe:
            logging.error(f"Erro de permissão ao acessar {file_path}: {pe}")
            error_type = self.error_handler.handle_extraction_error(pe, {"file": file_path})
            recovery_action = self.error_handler.suggest_recovery_action(error_type, file_path)
            logging.info(f"Sugestão de recuperação: {recovery_action['action']}")
            return False
            
        except MemoryError as me:
            logging.critical(f"Erro de memória ao processar {file_path}: {me}")
            messagebox.showerror("Erro de Memória", 
                            f"O arquivo {file_name} é muito grande para ser processado. Considere dividir o documento.")
            return False
            
        except Exception as e:
            logging.error(f"Erro desconhecido ao processar {file_path}: {e}")
            error_type = self.error_handler.handle_extraction_error(e, {"file": file_path})
            recovery_action = self.error_handler.suggest_recovery_action(error_type, file_path)
            logging.info(f"Sugestão de recuperação: {recovery_action['action']}")
            return False
    
    def _generate_summary_and_toc(self, docx_path: str, text: str) -> bool:
        """
        Gera sumário e tabela de conteúdo para o documento DOCX
        
        Args:
            docx_path (str): Caminho do documento DOCX
            text (str): Texto extraído do documento
        
        Returns:
            bool: True se bem-sucedido, False caso contrário
        """
        try:
            # Verificar se temos API key do Mistral (necessário para a geração)
            if not hasattr(self, 'mistral_ocr') or not self.mistral_ocr.api_key:
                logging.warning("API Key do Mistral não configurada. Não é possível gerar sumário.")
                messagebox.showwarning("Aviso", "API Key do Mistral é necessária para gerar sumário. Por favor, configure uma chave válida.")
                return False
            
            logging.info(f"Iniciando geração de sumário e tabela de conteúdo para {os.path.basename(docx_path)}")
            
            # Criar o enhancer com a API key do Mistral
            enhancer = DocumentEnhancer(self.mistral_ocr.api_key)
            
            # Processar o documento para adicionar sumário
            success = enhancer.process_document(docx_path, "mistral-large-latest")
            
            if success:
                logging.info(f"Sumário e tabela de conteúdo gerados com sucesso para {os.path.basename(docx_path)}")
                return True
            else:
                logging.error(f"Falha ao gerar sumário e tabela de conteúdo para {os.path.basename(docx_path)}")
                return False
                
        except Exception as e:
            logging.error(f"Erro ao gerar sumário e tabela de conteúdo: {e}")
            # Exibir erro apenas para debugging, não para o usuário final
            print(f"Erro na geração de sumário: {e}")
            return False

    def _generate_docx(self, file_path: str, output_dir: str, paragraphs: List[Tuple[str, str]]):
        """Gera documento DOCX formatado"""
        try:
            docx_path = os.path.join(output_dir,
                              f"{os.path.splitext(os.path.basename(file_path))[0]}.docx")
            
            # Criar documento
            doc = Document()
            
            # Configurar estilos de parágrafo
            DocxFormatter.setup_document_styles(doc)
            
            # Adicionar título do documento - sanitizando o nome do arquivo
            safe_filename = DocxFormatter.sanitize_text_for_xml(os.path.basename(file_path))
            doc.add_heading(f"Documento: {safe_filename}", level=0)
            
            # Adicionar metadados de processamento
            doc.add_paragraph(f"Processado com: {self.ocr_type_var.get().capitalize()} OCR")
            doc.add_paragraph(f"Data de processamento: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            doc.add_paragraph(f"Idioma: {self.lang_var.get()}")
            doc.add_paragraph("").paragraph_format.space_after = Pt(20)  # Espaço extra
            
            # Adicionar parágrafos formatados
            success_count = 0
            total_paragraphs = len(paragraphs)
            
            for para_text, para_type in paragraphs:
                try:
                    DocxFormatter.add_paragraph_with_style(doc, para_text, para_type)
                    success_count += 1
                except Exception as e:
                    # Registra o erro, mas continua processando os próximos parágrafos
                    logging.warning(f"Erro ao processar parágrafo: {str(e)[:100]}...")
            
            # Salvar documento
            doc.save(docx_path)
            logging.info(f"Documento DOCX criado: {docx_path} ({success_count}/{total_paragraphs} parágrafos processados)")
            
        except Exception as e:
            logging.error(f"Erro ao gerar DOCX para {file_path}: {e}")
            # Não propaga a exceção para permitir que o processamento continue com outros arquivos
            # Em vez de usar "raise", retornamos False para indicar falha
            return False
        
        return True
    


    def _generate_json(self, file_path: str, output_dir: str, text: str, paragraphs: List[Tuple[str, str]]):
        """Gera arquivo JSONL com entradas validadas"""
        try:
            entry = JsonFormatter.create_mistral_entry(text, paragraphs)
            
            if not entry:  # Entrada vazia ou sem resposta
                logging.warning(f"Ignorando entrada inválida para {file_path}")
                return False
                
            if not validate_mistral_entry(entry):
                logging.warning(f"Formato de entrada inválido para {file_path}")
                return False
                
            output_file = os.path.join(output_dir, "mistral_dataset.jsonl")
            
            # Use o lock global para acesso exclusivo ao arquivo
            with write_lock:  # Agora write_lock está definido
                with open(output_file, 'a', encoding='utf-8') as f:
                    f.write(json.dumps(entry, ensure_ascii=False) + '\n')
                
            return True
        except Exception as e:
            logging.error(f"Erro JSON: {e}")
            return False
        
    def _update_api_stats(self):
        """Atualiza estatísticas de uso da API na interface"""
        # Verificação defensiva para evitar o erro
        if hasattr(self, 'active_requests_label') and hasattr(self, 'mistral_ocr'):
            self.active_requests_label.config(text=str(self.mistral_ocr.active_requests))
        
        if hasattr(self, 'tokens_used_label') and hasattr(self, 'mistral_ocr'):
            self.tokens_used_label.config(text=str(self.mistral_ocr.total_tokens_used))
        
        if hasattr(self, 'total_calls_label') and hasattr(self, 'mistral_ocr'):
            self.total_calls_label.config(text=str(self.mistral_ocr.api_calls_count))
        
        # Atualiza status de conexão se API key estiver configurada
        if hasattr(self, 'api_connection_label') and hasattr(self, 'mistral_ocr'):
            if self.mistral_ocr.api_key:
                self.api_connection_label.config(text="Conectado", foreground="green")
            else:
                self.api_connection_label.config(text="Desconectado", foreground="red")
        
        # Programar próxima atualização (a cada 2 segundos)
        self.after(2000, self._update_api_stats)

class DocumentDataExtractor:
    """
    Extrator avançado de dados estruturados para diferentes tipos de documentos
    """
    def __init__(self, mistral_api_key: Optional[str] = None):
        """
        Inicializa o extrator de dados
        
        Args:
            mistral_api_key (str, optional): Chave API da Mistral para processamento avançado
        """
        self.mistral_client = Mistral(api_key=mistral_api_key) if mistral_api_key else None
        
        # Dicionário de padrões de extração por tipo de documento
        self.extraction_patterns = {
            'juridico': {
                'tipo_documento': r'\b(PROCESSO|PETIÇÃO|RECURSO|AÇÃO)\b',
                'numero_processo': r'\b(?:PROCESSO|PROTOCOLO)\s*[Nº]?\s*(\d{4,20})\b',
                'data_documento': r'\b(\d{1,2}/\d{1,2}/\d{2,4})\b',
                'valor_causa': r'\bVALOR\s*(?:DA\s*CAUSA)?\s*[R$]?\s*(\d+(?:\.\d{3})*,\d{2})\b'
            },
            'fiscal': {
                'tipo_documento': r'\b(NOTA FISCAL|CUPOM FISCAL|NFe)\b',
                'numero_documento': r'\b(NFe|Nota Fiscal)\s*[Nº]?\s*(\d{6,12})\b',
                'cnpj': r'\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b',
                'data_emissao': r'\b(\d{1,2}/\d{1,2}/\d{2,4})\b',
                'valor_total': r'\bVALOR\s*(?:TOTAL)?\s*[R$]?\s*(\d+(?:\.\d{3})*,\d{2})\b',
                'tipo_pagamento': r'\b(BOLETO|PIX|TRANSFERÊNCIA|CARTÃO)\b'
            },
            'bancario': {
                'tipo_documento': r'\b(EXTRATO|COMPROVANTE|CONTRACHEQUE)\b',
                'conta': r'\bCONTA\s*[Nº]?\s*(\d{6,10})\b',
                'agencia': r'\bAGÊNCIA\s*[Nº]?\s*(\d{4})\b',
                'data_lancamento': r'\b(\d{1,2}/\d{1,2}/\d{2,4})\b',
                'valor_lancamento': r'\bVALOR\s*[R$]?\s*(\d+(?:\.\d{3})*,\d{2})\b',
                'tipo_lancamento': r'\b(CRÉDITO|DÉBITO|TRANSFERÊNCIA|PAGAMENTO)\b'
            }
        }
    
    def _save_api_key(self):
        """Salva a chave da API em arquivo"""
        api_key = self.api_key_entry.get().strip()
        if not api_key:
            messagebox.showwarning("Aviso", "A chave da API não pode estar vazia.")
            return
        
        success = self._save_api_key_to_file(api_key)
        if success:
            messagebox.showinfo("Sucesso", "Chave da API salva com sucesso no arquivo.")
            self.api_status_label.config(
                text="Status: API key salva em arquivo com sucesso",
                foreground="green"
            )
        else:
            messagebox.showerror("Erro", "Não foi possível salvar a chave da API.")

    def _save_api_key_to_file(self, api_key: str) -> bool:
        """Salva a chave da API em um arquivo criptografado"""
        try:
            # Diretório de configuração na pasta do usuário
            config_dir = os.path.join(os.path.expanduser("~"), ".ocr_processor")
            os.makedirs(config_dir, exist_ok=True)
            
            # Caminho do arquivo de chave
            key_file = os.path.join(config_dir, "api_key.dat")
            
            # Criptografia simples (versão mais segura)
            # Em ambiente de produção, usar biblioteca especializada como cryptography
            import base64
            
            try:
                # Tenta usar o módulo secrets (Python 3.6+)
                import secrets
                salt = secrets.token_bytes(16)
            except ImportError:
                # Alternativa para Python < 3.6
                import os
                import random
                # Criar 16 bytes aleatórios como salt
                salt = bytes([random.randint(0, 255) for _ in range(16)])
            
            # Combinar salt com a chave e codificar
            salted_key = salt + api_key.encode('utf-8')
            encoded_key = base64.b64encode(salted_key)
            
            # Salvar com o salt no início para poder descriptografar depois
            with open(key_file, 'wb') as f:
                f.write(encoded_key)
                
            logging.info("Chave da API salva com sucesso")
            return True
        
        except Exception as e:
            logging.error(f"Erro ao salvar chave da API: {e}")
            return False
  
    def normalize_value(self, value: str) -> float:
        """
        Normaliza valores monetários
        
        Args:
            value (str): Valor monetário em formato de string
        
        Returns:
            float: Valor normalizado
        """
        if not value:
            return 0.0
        
        # Remove formatação
        value = value.replace('R$', '').replace('.', '').replace(',', '.')
        
        try:
            return float(value)
        except ValueError:
            return 0.0
    
    def normalize_date(self, date_str: str) -> Optional[datetime]:
        """
        Normaliza datas em diferentes formatos
        
        Args:
            date_str (str): Data em formato de string
        
        Returns:
            Optional[datetime]: Data normalizada ou None
        """
        if not date_str:
            return None
        
        # Formatos de data para tentar
        date_formats = [
            '%d/%m/%Y',   # 31/12/2023
            '%d/%m/%y',   # 31/12/23
            '%m/%d/%Y',   # 12/31/2023
            '%Y-%m-%d'    # 2023-12-31
        ]
        
        for fmt in date_formats:
            try:
                return datetime.strptime(date_str.strip(), fmt)
            except ValueError:
                continue
        
        return None
    
    def extract_document_data(self, text: str, documento_tipo: Optional[str] = None) -> Dict[str, Any]:
        """
        Extrai dados estruturados de um documento
        
        Args:
            text (str): Texto completo do documento
            documento_tipo (str, optional): Tipo de documento (se conhecido)
        
        Returns:
            Dict[str, Any]: Dados estruturados extraídos
        """
        # Se o tipo não for fornecido, tenta identificar
        if not documento_tipo:
            for tipo, patterns in self.extraction_patterns.items():
                if re.search(patterns['tipo_documento'], text, re.IGNORECASE):
                    documento_tipo = tipo
                    break
        
        # Se ainda não identificou, retorna dados básicos
        if not documento_tipo:
            return {
                'tipo_documento': 'Não identificado',
                'texto_original': text[:500]  # Primeiros 500 caracteres
            }
        
        # Extrai dados baseado no tipo de documento
        dados_extraidos = {
            'tipo_documento': documento_tipo.capitalize()
        }
        
        # Padrões para o tipo de documento específico
        patterns = self.extraction_patterns.get(documento_tipo, {})
        
        # Extração de dados
        for campo, padrao in patterns.items():
            if campo == 'tipo_documento':
                continue
            
            # Busca correspondências
            match = re.search(padrao, text, re.IGNORECASE)
            if match:
                # Extrai o valor da correspondência
                valor = match.group(1) if match.groups() else match.group(0)
                
                # Normaliza baseado no nome do campo
                if 'valor' in campo or 'total' in campo:
                    dados_extraidos[campo] = self.normalize_value(valor)
                elif 'data' in campo:
                    dados_extraidos[campo] = self.normalize_date(valor)
                else:
                    dados_extraidos[campo] = valor
        
        # Processamento adicional com IA (se API disponível)
        if self.mistral_client:
            dados_extraidos = self._enrich_with_ai(text, dados_extraidos)
        
        return dados_extraidos
    
    def _enrich_with_ai(self, text: str, dados_extraidos: Dict[str, Any]) -> Dict[str, Any]:
        """
        Enriquece os dados extraídos usando IA
        
        Args:
            text (str): Texto completo do documento
            dados_extraidos (Dict[str, Any]): Dados já extraídos
        
        Returns:
            Dict[str, Any]: Dados enriquecidos
        """
        try:
            # Prompt para extração adicional
            messages = [
                {
                    "role": "system",
                    "content": """Você é um assistente especializado em extração de informações de documentos. 
                    Analise o texto fornecido e extraia informações adicionais não capturadas pelos padrões básicos.
                    Forneça dados em formato JSON."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Dados já extraídos: {json.dumps(dados_extraidos)}
                    
                    Texto do documento: {text[:2000]}
                    
                    Por favor, forneça informações adicionais relevantes em JSON.
                    Foque em campos não preenchidos que possam ser importantes.
                    """
                }
            ]
            
            # Chamada à API Mistral
            response = self.mistral_client.chat.complete(
                model="mistral-large-latest",
                messages=messages,
                response_format={"type": "json_object"}
            )
            
            # Mescla dados adicionais
            dados_ai = json.loads(response.choices[0].message.content)
            dados_extraidos.update(dados_ai)
        
        except Exception as e:
            logging.error(f"Erro no enriquecimento por IA: {e}")
        
        return dados_extraidos
    
    def process_document_batch(self, documentos: List[str]) -> pd.DataFrame:
        """
        Processa um lote de documentos
        
        Args:
            documentos (List[str]): Lista de caminhos de documentos
        
        Returns:
            pd.DataFrame: DataFrame com dados extraídos
        """
        # Lista para armazenar dados
        dados_documentos = []
        
        # Processa cada documento
        for documento in documentos:
            try:
                # Usa o método _extract_text da classe OCR
                texto = extract_text(documento)
                
                # Extrai dados
                dados = self.extract_document_data(texto)
                
                # Adiciona caminho do documento
                dados['caminho_documento'] = documento
                
                dados_documentos.append(dados)
            
            except Exception as e:
                logging.error(f"Erro ao processar documento {documento}: {e}")
        
        # Converte para DataFrame
        return pd.DataFrame(dados_documentos)

class DocumentEnhancer:
    """Classe para aprimorar documentos com sumário e tabela de conteúdo usando IA"""
    
    def __init__(self, api_key):
        """
        Inicializa o DocumentEnhancer com a chave da API Mistral
        
        Args:
            api_key (str): Chave de API da Mistral
        """
        self.client = Mistral(api_key=api_key)
    
    def generate_summary_and_toc(self, text, model="mistral-large-latest"):
        """
        Gera sumário e tabela de conteúdo usando a API da Mistral
        
        Args:
            text (str): Texto completo do documento
            model (str, optional): Modelo de IA a ser usado
        
        Returns:
            dict: Dicionário contendo sumário e tabela de conteúdo
        """
        # Limita o texto para evitar excesso de tokens (15000 caracteres é um bom limite)
        truncated_text = text[:15000]
        
        messages = [
            {
                "role": "system", 
                "content": """Você é um assistente especializado em análise de documentos. 
                Seu objetivo é gerar um sumário conciso e uma tabela de conteúdo detalhada 
                para o documento fornecido. Siga estas diretrizes:
                
                1. Sumário Executivo:
                - Máximo de 3-5 parágrafos
                - Capture a essência do documento
                - Destaque os pontos-chave
                
                2. Tabela de Conteúdo:
                - Identifique seções principais e subseções
                - Use numeração hierárquica (1, 1.1, 1.2, etc.)
                - Forneça breve descrição de cada seção
                
                Responda em formato JSON com as seguintes chaves:
                - "summary": Sumário executivo em texto
                - "table_of_contents": Tabela de conteúdo detalhada
                """
            },
            {
                "role": "user", 
                "content": f"Gere um sumário e tabela de conteúdo para o seguinte documento:\n\n{truncated_text}"
            }
        ]
        
        try:
            response = self.client.chat.complete(
                model=model,
                messages=messages,
                response_format={"type": "json_object"}
            )
            
            # Parse do JSON retornado
            result = json.loads(response.choices[0].message.content)
            return result
        except Exception as e:
            logging.error(f"Erro ao gerar sumário: {e}")
            return None
    
    def add_summary_and_toc_to_docx(self, docx_path, summary_data):
        """
        Adiciona sumário e tabela de conteúdo no início do documento DOCX
        
        Args:
            docx_path (str): Caminho para o arquivo DOCX
            summary_data (dict): Dados de sumário gerados pela API
        """
        try:
            # Abre o documento existente
            doc = Document(docx_path)
            
            # Cria uma cópia temporária do conteúdo existente
            old_content = []
            for paragraph in doc.paragraphs:
                old_content.append(paragraph.text)
            
            # Limpa o documento
            for i in range(len(doc.paragraphs)-1, -1, -1):
                p = doc.paragraphs[i]
                p._element.getparent().remove(p._element)
            
            # Insere sumário executivo
            doc.add_heading('Sumário Executivo', level=1)
            summary_para = doc.add_paragraph(summary_data.get('summary', 'Sumário não disponível'))
            summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Insere uma quebra de página
            doc.add_page_break()
            
            # Insere tabela de conteúdo
            doc.add_heading('Tabela de Conteúdo', level=1)
            toc_para = doc.add_paragraph(summary_data.get('table_of_contents', 'Tabela de conteúdo não disponível'))
            toc_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Insere uma quebra de página
            doc.add_page_break()
            
            # Adiciona o conteúdo original de volta
            for text in old_content:
                doc.add_paragraph(text)
            
            # Salva o documento atualizado
            doc.save(docx_path)
            logging.info(f"Sumário e tabela de conteúdo adicionados a {docx_path}")
            return True
        
        except Exception as e:
            logging.error(f"Erro ao adicionar sumário ao documento: {e}")
            return False
    
    def process_document(self, docx_path, model="mistral-large-latest"):
        """
        Processa um documento para adicionar sumário e tabela de conteúdo
        
        Args:
            docx_path (str): Caminho para o arquivo DOCX
            model (str, optional): Modelo de IA a ser usado
        
        Returns:
            bool: True se processamento for bem-sucedido, False caso contrário
        """
        try:
            # Extrai texto do documento
            doc = Document(docx_path)
            full_text = '\n'.join([para.text for para in doc.paragraphs if para.text])
            
            # Gera sumário e tabela de conteúdo
            summary_data = self.generate_summary_and_toc(full_text, model)
            
            # Adiciona sumário e tabela de conteúdo ao documento
            if summary_data:
                return self.add_summary_and_toc_to_docx(docx_path, summary_data)
            return False
        
        except Exception as e:
            logging.error(f"Erro no processamento do documento: {e}")
            return False

from pydantic import BaseModel, ValidationError

class MistralMessage(BaseModel):
    role: str
    content: str

class MistralEntry(BaseModel):
    messages: List[MistralMessage]

def validate_mistral_entry(entry: Dict) -> bool:
    """Validação rigorosa do formato Mistral"""
    if not isinstance(entry.get("messages"), list):
        return False
        
    if len(entry["messages"]) < 2:
        return False
        
    # Primeira mensagem deve ser do usuário
    if entry["messages"][0]["role"] != "user":
        return False
        
    # Última mensagem deve ser do assistente
    if entry["messages"][-1]["role"] != "assistant":
        return False
        
    # Não permite mensagens consecutivas do mesmo papel
    last_role = None
    for msg in entry["messages"]:
        if msg["role"] == last_role:
            return False
        last_role = msg["role"]
            
    return True

from enum import Enum, auto

class DocumentErrorType(Enum):
    """
    Categorização detalhada de tipos de erros de documento
    """
    OCR_LOW_QUALITY = auto()
    INCOMPLETE_DOCUMENT = auto()
    UNSUPPORTED_FORMAT = auto()
    METADATA_EXTRACTION_FAILURE = auto()
    NETWORK_ERROR = auto()
    VALIDATION_ERROR = auto()

class DocumentErrorHandler:
    """
    Manipulador avançado de erros de processamento de documentos
    """
    def __init__(self):
        # Configuração de logging
        self.logger = logging.getLogger(__name__)
    
    def handle_extraction_error(self, error: Exception, context: Dict[str, str]) -> DocumentErrorType:
        """
        Classifica e trata erros de extração
        
        Args:
            error (Exception): Erro capturado
            context (Dict[str, str]): Contexto da extração
        
        Returns:
            DocumentErrorType: Tipo de erro classificado
        """
        error_message = str(error)
        error_trace = traceback.format_exc()
        
        # Classificação de erros
        if "low resolution" in error_message.lower() or "insufficient quality" in error_message.lower():
            self.logger.error(f"Erro de baixa resolução: {error_message}")
            return DocumentErrorType.OCR_LOW_QUALITY
        
        elif "network" in error_message.lower() or "connection" in error_message.lower():
            self.logger.error(f"Erro de rede: {error_message}")
            return DocumentErrorType.NETWORK_ERROR
        
        elif "unsupported format" in error_message.lower() or "not a PDF" in error_message.lower():
            self.logger.error(f"Formato não suportado: {error_message}")
            return DocumentErrorType.UNSUPPORTED_FORMAT
        
        elif "incomplete" in error_message.lower() or "corrupt" in error_message.lower():
            self.logger.error(f"Documento incompleto: {error_message}")
            return DocumentErrorType.INCOMPLETE_DOCUMENT
        
        elif "validation" in error_message.lower() or "invalid data" in error_message.lower():
            self.logger.error(f"Erro de validação: {error_message}")
            return DocumentErrorType.VALIDATION_ERROR
        
        # Log do erro detalhado
        self.logger.error(
            f"Erro não classificado:\n"
            f"Mensagem: {error_message}\n"
            f"Contexto: {context}\n"
            f"Rastreamento: {error_trace}"
        )
        
        return DocumentErrorType.METADATA_EXTRACTION_FAILURE
    
    def suggest_recovery_action(self, error_type: DocumentErrorType, document_path: str) -> Dict[str, Any]:
        """
        Sugere ações de recuperação baseadas no tipo de erro
        
        Args:
            error_type (DocumentErrorType): Tipo de erro
            document_path (str): Caminho do documento
        
        Returns:
            Dict[str, str]: Sugestões de recuperação
        """
        recovery_actions = {
            DocumentErrorType.OCR_LOW_QUALITY: {
                'action': 'Reprocessar com pré-processamento de imagem',
                'steps': [
                    'Aumentar contraste',
                    'Aplicar filtro de nitidez',
                    'Remover ruído'
                ]
            },
            DocumentErrorType.NETWORK_ERROR: {
                'action': 'Tentar novamente',
                'steps': [
                    'Verificar conexão de internet',
                    'Aguardar e tentar novamente',
                    'Verificar serviços online'
                ]
            },
            DocumentErrorType.UNSUPPORTED_FORMAT: {
                'action': 'Converter documento',
                'steps': [
                    'Converter para PDF',
                    'Verificar formato de origem',
                    'Usar ferramentas de conversão'
                ]
            },
            DocumentErrorType.INCOMPLETE_DOCUMENT: {
                'action': 'Verificar documento',
                'steps': [
                    'Verificar se o documento está corrompido',
                    'Tentar obter uma cópia completa',
                    'Usar software de reparo de PDF'
                ]
            },
            DocumentErrorType.METADATA_EXTRACTION_FAILURE: {
                'action': 'Revisão manual',
                'steps': [
                    'Verificar documento original',
                    'Realizar extração manual',
                    'Documentar problema'
                ]
            },
            DocumentErrorType.VALIDATION_ERROR: {
                'action': 'Corrigir dados',
                'steps': [
                    'Verificar formato dos dados',
                    'Corrigir inconsistências',
                    'Ajustar parâmetros de validação'
                ]
            }
        }
        
        # Registra sugestão
        self.logger.info(
            f"Sugestão de recuperação para {error_type}: "
            f"{recovery_actions.get(error_type, 'Sem ação definida')}"
        )
        
        return recovery_actions.get(error_type, {
            'action': 'Análise manual necessária',
            'steps': ['Verificar o documento manualmente']
        })

if __name__ == "__main__":
    # Verifica Tesseract OCR
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
    except EnvironmentError:
        if 'tkinter' in sys.modules:
            root = tk.Tk()
            root.withdraw()
            messagebox.showerror("Erro", "Tesseract OCR não encontrado! Instale em https://github.com/UB-Mannheim/tesseract/wiki")
            root.destroy()
        else:
            print("Erro: Tesseract OCR não encontrado! Instale em https://github.com/UB-Mannheim/tesseract/wiki")
        exit(1)
    
    # Verifica Poppler no início da aplicação
    if not check_poppler_installed():
        show_poppler_instructions()
        # Continua a execução mas avisa o usuário
    
    # Inicia a aplicação
    PDFProcessorApp().mainloop()
