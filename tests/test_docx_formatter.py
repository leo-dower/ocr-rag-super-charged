
import unittest
from docx import Document
from src.utils.docx_formatter import DocxFormatter

class TestDocxFormatter(unittest.TestCase):

    def test_sanitize_text_for_xml(self):
        self.assertEqual(DocxFormatter.sanitize_text_for_xml("Test\x00\x0b\x0c\x0e\x1fString"), "TestString")

    def test_setup_document_styles(self):
        doc = Document()
        DocxFormatter.setup_document_styles(doc)
        self.assertIn('Normal Paragraph', doc.styles)
        self.assertIn('Artigo', doc.styles)
        self.assertIn('Destaque', doc.styles)

    def test_add_paragraph_with_style(self):
        doc = Document()
        DocxFormatter.setup_document_styles(doc)
        DocxFormatter.add_paragraph_with_style(doc, "<p>Test</p>", "normal")
        self.assertIn("Test", doc.paragraphs[-1].text)

if __name__ == '__main__':
    unittest.main()
