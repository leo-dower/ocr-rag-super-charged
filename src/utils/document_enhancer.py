

import json
from mistralai import Mistral
import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocumentEnhancer:
    """Classe para aprimorar documentos com sumário e tabela de conteúdo usando IA"""
    
    def __init__(self, api_key):
        self.client = Mistral(api_key=api_key)
    
    def generate_summary_and_toc(self, text, model="mistral-large-latest"):
        truncated_text = text[:15000]
        
        messages = [
            {
                "role": "system", 
                "content": "Você é um assistente especializado em análise de documentos. Seu objetivo é gerar um sumário conciso e uma tabela de conteúdo detalhada para o documento fornecido. Siga estas diretrizes:\n\n1. Sumário Executivo:\n- Máximo de 3-5 parágrafos\n- Capture a essência do documento\n- Destaque os pontos-chave\n\n2. Tabela de Conteúdo:\n- Identifique seções principais e subseções\n- Use numeração hierárquica (1, 1.1, 1.2, etc.)\n- Forneça breve descrição de cada seção\n\nResponda em formato JSON com as seguintes chaves:\n- \"summary\": Sumário executivo em texto\n- \"table_of_contents\": Tabela de conteúdo detalhada"
            },
            {
                "role": "user", 
                "content": f"Gere um sumário e tabela de conteúdo para o seguinte documento:\n\n{truncated_text}"
            }
        ]
        
        try:
            response = self.client.chat.complete(
                model=model,
                messages=messages,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            return result
        except Exception as e:
            logging.error(f"Erro ao gerar sumário: {e}")
            return None
    
    def add_summary_and_toc_to_docx(self, docx_path, summary_data):
        try:
            doc = Document(docx_path)
            
            old_content = []
            for paragraph in doc.paragraphs:
                old_content.append(paragraph.text)
            
            for i in range(len(doc.paragraphs)-1, -1, -1):
                p = doc.paragraphs[i]
                p._element.getparent().remove(p._element)
            
            doc.add_heading('Sumário Executivo', level=1)
            summary_para = doc.add_paragraph(summary_data.get('summary', 'Sumário não disponível'))
            summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            doc.add_page_break()
            
            doc.add_heading('Tabela de Conteúdo', level=1)
            toc_para = doc.add_paragraph(summary_data.get('table_of_contents', 'Tabela de conteúdo não disponível'))
            toc_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            doc.add_page_break()
            
            for para_text in old_content:
                doc.add_paragraph(para_text)

            doc.save(docx_path)
            return True
        except Exception as e:
            logging.error(f"Erro ao adicionar sumário ao DOCX: {e}")
            return False

