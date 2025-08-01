
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import logging
from html import unescape

class DocxFormatter:
    """Classe para formatação avançada de documentos DOCX"""
    
    @staticmethod
    def sanitize_text_for_xml(text: str) -> str:
        """
        Limpa o texto de caracteres incompatíveis com XML
        
        Args:
            text: O texto a ser limpo
            
        Returns:
            str: Texto limpo compatível com XML
        """
        if not text:
            return ""
            
        def is_xml_char(c):
            cp = ord(c)
            return (
                cp == 0x9 or
                cp == 0xA or
                cp == 0xD or
                (0x20 <= cp <= 0xD7FF) or
                (0xE000 <= cp <= 0xFFFD) or
                (0x10000 <= cp <= 0x10FFFF)
            )
        
        return ''.join(c for c in text if is_xml_char(c))
    
    @staticmethod
    def setup_document_styles(doc: Document) -> None:
        """Configura estilos do documento para melhor formatação de parágrafos"""
        if 'Normal Paragraph' not in doc.styles:
            normal_style = doc.styles.add_style('Normal Paragraph', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            normal_style.paragraph_format.first_line_indent = Inches(0.5)
            normal_style.paragraph_format.space_after = Pt(10)
        
        if 'Artigo' not in doc.styles:
            artigo_style = doc.styles.add_style('Artigo', WD_STYLE_TYPE.PARAGRAPH)
            artigo_style.font.size = Pt(12)
            artigo_style.font.bold = True
            artigo_style.paragraph_format.space_before = Pt(12)
            artigo_style.paragraph_format.space_after = Pt(6)
        
        if 'Destaque' not in doc.styles:
            destaque_style = doc.styles.add_style('Destaque', WD_STYLE_TYPE.PARAGRAPH)
            destaque_style.font.size = Pt(12)
            destaque_style.font.bold = True
            destaque_style.paragraph_format.first_line_indent = Inches(0.5)
            destaque_style.paragraph_format.space_after = Pt(10)
    
    @staticmethod
    def add_paragraph_with_style(doc: Document, text: str, para_type: str) -> None:
        """Adiciona parágrafo com o estilo apropriado baseado no tipo de parágrafo"""
        text = unescape(text)
        clean_text = text.replace('**', '')
        clean_text = DocxFormatter.sanitize_text_for_xml(clean_text)
        
        if not clean_text:
            return
        
        try:
            if para_type == "titulo":
                doc.add_heading(clean_text, level=1)
            elif para_type == "artigo":
                p = doc.add_paragraph(clean_text, style='Artigo')
            elif para_type == "destaque":
                p = doc.add_paragraph(clean_text, style='Destaque')
            else:
                p = doc.add_paragraph(clean_text, style='Normal Paragraph')
        except ValueError as e:
            logging.warning(f"Não foi possível adicionar um parágrafo: {e}. Texto: {clean_text[:50]}...")
